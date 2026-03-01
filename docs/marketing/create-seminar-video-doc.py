#!/usr/bin/env python3
"""Generate Word document with Muay Thai seminar promo video scenarios."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Styles setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x21, 0x25, 0x29)

style_h1 = doc.styles['Heading 1']
style_h1.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)
style_h1.font.size = Pt(22)
style_h1.font.bold = True

style_h2 = doc.styles['Heading 2']
style_h2.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)
style_h2.font.size = Pt(16)
style_h2.font.bold = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def add_orange_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
    return p


# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RIGHT TRACK \u00d7 MUAY THAI WARRIORS')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Promo Video Scenarios')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Seminar: "Train Hard. Fight Smart. Stay Injury-Free."')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nFebruary 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

doc.add_page_break()

# ============================================================
# OVERVIEW
# ============================================================
doc.add_heading('Overview', level=1)

doc.add_paragraph(
    'This document contains ready-to-shoot video scenarios for promoting the '
    '"Train Hard. Fight Smart. Stay Injury-Free." injury prevention seminar \u2014 '
    'a collaboration between Right Track Physiotherapy and Muay Thai Warriors gym.'
)

doc.add_heading('Seminar Details', level=2)

add_table(doc,
    ['Detail', 'Info'],
    [
        ['Event', 'Injury Prevention Seminar for Fighters'],
        ['Duration', '90 minutes'],
        ['Price', '\u20ac25 per person'],
        ['Capacity', '30 participants'],
        ['Venue', 'Muay Thai Warriors Gym, Nicosia'],
        ['Speaker', 'Antonis Petri, BSc, OMPT \u2014 Physiotherapist'],
        ['Date', 'April 4, 2026'],
        ['Exclusive offer', 'Fighter Recovery Package \u2014 \u20ac89 (only for attendees)'],
    ]
)

doc.add_paragraph('')
doc.add_heading('Shoot Details', level=2)

add_table(doc,
    ['Detail', 'Info'],
    [
        ['Shoot location', 'Muay Thai Warriors Gym'],
        ['People on camera', 'Antonis Petri + Gym Owner / Head Coach (optional: Antonis solo)'],
        ['Format', 'Vertical 9:16 (1080\u00d71920) \u2014 Instagram Reels / Stories / TikTok / Shorts'],
        ['Target duration', '30\u201360 seconds (edited)'],
        ['Audio', 'Lapel mic on Antonis (mandatory \u2014 gyms are loud)'],
        ['Language', 'Greek (primary) \u2014 subtitles added in post-production'],
    ]
)

doc.add_page_break()

# ============================================================
# SCENARIO A
# ============================================================
doc.add_heading('Scenario A: "Two Worlds"', level=1)

p = doc.add_paragraph()
run = p.add_run('Antonis + Coach  |  45\u201360 seconds  |  Best option if coach is available')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)

doc.add_paragraph(
    'Concept: Two worlds \u2014 the fight gym and physiotherapy \u2014 come together. '
    'The contrast between raw combat training and professional injury prevention '
    'creates a compelling narrative that speaks to both audiences.'
)

doc.add_heading('Shot List', level=2)

add_table(doc,
    ['Sec', 'Shot', 'Who', 'Audio'],
    [
        ['0\u20133', 'Coach hitting pads / heavy bag (b-roll, close-up of impact)', 'Coach', 'Impact sounds (ambient)'],
        ['3\u20138', 'Coach to camera: "My fighters train hard. But injuries are part of the game"', 'Coach', 'Speech'],
        ['8\u201313', 'Antonis walks into the gym, handshake with the coach (wide shot)', 'Both', 'Music builds'],
        ['13\u201320', 'Antonis to camera: "I\'m Antonis, a physiotherapist. I work with fighters, footballers, and athletes. And we decided to do something about it"', 'Antonis', 'Speech'],
        ['20\u201330', 'Both standing side by side.\nCoach: "We\'re hosting a seminar right here, in our gym."\nAntonis: "90 minutes. Warm-up, joint protection, recovery \u2014 everything a fighter needs to train without injuries"', 'Both', 'Dialogue'],
        ['30\u201340', 'B-roll montage: bags, ring, fighters training (if present), club logo on the wall', '\u2014', 'Music'],
        ['40\u201350', 'Antonis to camera: "30 spots. \u20ac25. Link in the description"', 'Antonis', 'Speech'],
        ['50\u201355', 'End screen with text: April 4, venue, price, registration link, both logos', '\u2014', 'Music'],
    ]
)

doc.add_paragraph('')
add_orange_heading(doc, 'Why it works')
doc.add_paragraph(
    'Shows a real partnership. The coach brings trust from his own audience. '
    'Two faces = two communities reached. The gym environment provides authentic context '
    'that no studio can replicate.'
)

add_orange_heading(doc, 'Director\'s notes')
for b in [
    'Open with action (punching) \u2014 grabs attention in the first second',
    'The handshake moment is key \u2014 it visually communicates partnership',
    'Both should face the camera directly, not each other',
    'Shoot the coach\'s lines first (he may be less comfortable on camera \u2014 get it done while energy is high)',
    'Capture extra b-roll of the gym during any downtime',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# SCENARIO B
# ============================================================
doc.add_heading('Scenario B: "The Coach\'s Question"', level=1)

p = doc.add_paragraph()
run = p.add_run('Antonis + Coach  |  30\u201345 seconds  |  Fastest-paced option, strongest hook')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)

doc.add_paragraph(
    'Concept: Opens with a question that stops the scroll. Quick back-and-forth '
    'dialogue between the coach and Antonis creates a dynamic, engaging format. '
    'Follows the proven "problem \u2192 expert \u2192 solution" Reel formula.'
)

doc.add_heading('Shot List', level=2)

add_table(doc,
    ['Sec', 'Shot', 'Audio'],
    [
        ['0\u20133', 'Coach to camera, gym in background (medium shot)', 'Coach: "Do you know what the most common injury in Muay Thai is?"'],
        ['3\u20135', 'Quick cuts: kicks, sparring, pad work (b-roll, 3\u20134 clips)', 'Impact sounds'],
        ['5\u201312', 'Antonis to camera (same location, in the gym): "Ankle ligament damage. And 80% of fighters just tape it up and keep going. That\'s a mistake"', 'Antonis speaking'],
        ['12\u201318', 'Coach to camera: "That\'s why we brought Antonis in. He knows how to train hard without breaking down"', 'Coach speaking'],
        ['18\u201325', 'Antonis: "A seminar at Muay Thai Warriors. 90 minutes of practical knowledge. Warm-up, joints, recovery. Everything you can apply at your next training session"', 'Antonis speaking'],
        ['25\u201330', 'Both together, gesture to camera. Text overlay: April 4, venue, price, "Link in bio"', 'Music'],
    ]
)

doc.add_paragraph('')
add_orange_heading(doc, 'Why it works')
doc.add_paragraph(
    'The opening question is a pattern interrupt \u2014 it stops the scroll immediately. '
    'The fast pace (under 30 seconds of talking) matches how Reels are consumed. '
    'The "problem \u2192 expert \u2192 solution" structure is the highest-performing Reel formula.'
)

add_orange_heading(doc, 'Director\'s notes')
for b in [
    'The hook (first 3 seconds) makes or breaks this video \u2014 shoot it 3\u20134 times',
    'Coach and Antonis should stand in the SAME spot (continuity) \u2014 just swap who\'s in front',
    'Keep energy high \u2014 short sentences, confident delivery',
    'The stat ("80% of fighters...") adds credibility \u2014 Antonis should deliver it with conviction',
    'End card should have clear CTA with registration link',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# SCENARIO C
# ============================================================
doc.add_heading('Scenario C: "Antonis Solo in the Gym"', level=1)

p = doc.add_paragraph()
run = p.add_run('Antonis alone  |  30\u201340 seconds  |  Backup if coach is unavailable')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)

doc.add_paragraph(
    'Concept: Antonis walks through the gym, speaking directly to camera. '
    'The gym environment creates all the context needed. Simple format that can '
    'be shot in 10 minutes. Follows the "Injury Decoded" formula from the content strategy.'
)

doc.add_heading('Shot List', level=2)

add_table(doc,
    ['Sec', 'Shot', 'Audio'],
    [
        ['0\u20133', 'Antonis walks into the gym (camera follows from behind, then he turns)', 'Ambient gym sounds'],
        ['3\u20138', 'Antonis to camera, ring/bags in background: "This is Muay Thai Warriors. Fighters train here. And I\'m the physiotherapist who puts them back together"', 'Antonis speaking'],
        ['8\u201315', 'Antonis walks through the gym, touches a bag or gloves: "Soon, right here \u2014 an injury prevention seminar for fighters. 90 minutes of practical knowledge"', 'Antonis speaking'],
        ['15\u201325', 'Antonis stops, looks directly at camera: "A warm-up that actually works. Joint protection. Recovery without the myths. 30 spots, \u20ac25"', 'Antonis speaking'],
        ['25\u201335', '"If you train in any combat sport \u2014 this is for you. Link in bio" + end screen with text overlay', 'Antonis + Music'],
    ]
)

doc.add_paragraph('')
add_orange_heading(doc, 'Why it works')
doc.add_paragraph(
    'Minimal setup, maximum atmosphere. The "walk and talk" format feels natural '
    'and authentic. The gym does the heavy lifting visually \u2014 no need for graphics '
    'or effects. Antonis\'s confidence in a fight gym establishes him as someone '
    'who belongs in this world.'
)

add_orange_heading(doc, 'Director\'s notes')
for b in [
    'The walk-in shot is cinematic \u2014 use it to build anticipation',
    'Antonis should move slowly and deliberately through the gym (not rushing)',
    'Touch points (bag, gloves, ring ropes) create visual anchors',
    'Keep the camera slightly below eye level \u2014 makes the speaker look authoritative',
    'Can be shot entirely on a phone with a gimbal/stabilizer',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# SCENARIO D: STORIES
# ============================================================
doc.add_heading('Scenario D: Stories Series (Supplement)', level=1)

p = doc.add_paragraph()
run = p.add_run('3\u20135 Stories  |  Shot alongside any Reel scenario  |  Near-zero extra effort')
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)

doc.add_paragraph(
    'These Stories are shot on the same visit to the gym, alongside whichever '
    'Reel scenario you choose. They build anticipation, engage the audience with '
    'interactive elements, and drive traffic to the registration link.'
)

doc.add_heading('Story-by-Story Breakdown', level=2)

add_table(doc,
    ['#', 'Content', 'Format', 'Notes'],
    [
        ['1', 'Antonis selfie-video at the gym entrance: "Just arrived at Muay Thai Warriors, we\'re cooking something up for you..."', 'Selfie video, 10 sec', 'Casual, unscripted. Show the gym sign/entrance'],
        ['2', 'Pan shot of the gym interior \u2014 bags, ring, training area. Text overlay: "Something big is coming here soon"', 'Video + text overlay, 8 sec', 'No talking, just atmosphere. Add location tag'],
        ['3', 'Quick informal greeting with the coach \u2014 introduce each other, casual banter', 'Selfie video with both, 12 sec', 'Not scripted. Natural energy. Tag the gym account'],
        ['4', 'Poll sticker: "What\'s YOUR most common training injury?" Options: Knee / Shoulder / Ankle / Back', 'Photo or video + poll sticker', 'Engagement tool. Responses = content ideas'],
        ['5', '"All the answers \u2014 at our seminar. April 4th. Stay tuned" + countdown sticker or "Link in bio"', 'Text card or video, 8 sec', 'CTA. Add link sticker to registration page'],
    ]
)

doc.add_paragraph('')
add_orange_heading(doc, 'Why it works')
doc.add_paragraph(
    'Stories create a narrative arc: arrival \u2192 discovery \u2192 teaser \u2192 engagement \u2192 CTA. '
    'The poll sticker drives interaction (algorithm boost) and provides real data about '
    'the audience\'s pain points. Stories disappear in 24 hours, creating urgency.'
)

doc.add_page_break()

# ============================================================
# TECHNICAL CHECKLIST
# ============================================================
doc.add_heading('Technical Checklist for Shoot Day', level=1)

doc.add_heading('Equipment', level=2)
for item in [
    'Phone with good camera (iPhone 13+ or equivalent)',
    'Tripod or gimbal/stabilizer for smooth movement shots',
    'Lapel microphone for Antonis (CRITICAL \u2014 gyms are very loud)',
    'Fully charged battery + power bank',
    'Storage: minimum 5 GB free space',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Wardrobe', level=2)
for item in [
    'Antonis: dark t-shirt or polo (navy/black), no bright prints',
    'Right Track branded clothing if available',
    'Coach: his usual gym attire (authentic)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Shooting Tips', level=2)
for item in [
    'Shoot VERTICAL (9:16) for all Reels and Stories',
    'Also capture a few HORIZONTAL clips for the website and LinkedIn',
    'Shoot the hook line (first 3 seconds) at least 3 times \u2014 it\'s the most important part',
    'If the gym is dark, position near windows or turn on all available lights',
    'Capture 2\u20133 minutes of b-roll: bags, ring, gloves, club logo, training atmosphere',
    'Film at least one wide establishing shot of the gym exterior',
    'Have the coach speak naturally \u2014 imperfect delivery feels more authentic than polished scripts',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Post-Production', level=2)
for item in [
    'Add subtitles (white text, dark outline) \u2014 use CapCut or Captions app',
    'Add both logos (Right Track + Muay Thai Warriors) on the end screen',
    'Include text overlay with: April 4, venue, price (\u20ac25), registration link',
    'Background music: energetic but not overpowering (trending Reel audio if possible)',
    'Colour grading: warm tones, high contrast \u2014 not washed out',
    'Export at 1080\u00d71920 resolution, 30fps minimum',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')

# ============================================================
# RECOMMENDATION
# ============================================================
doc.add_heading('Recommendation', level=1)

p = doc.add_paragraph()
run = p.add_run('Best combination: Scenario A or B + Scenario D')
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph('')

add_table(doc,
    ['Situation', 'Shoot this', 'Total time'],
    [
        ['Coach is available and comfortable on camera', 'Scenario A (Two Worlds) + Scenario D (Stories)', '20\u201330 min'],
        ['Coach is available but limited time', 'Scenario B (The Question) + Scenario D (Stories)', '15\u201320 min'],
        ['Coach is unavailable', 'Scenario C (Solo) + Scenario D (Stories)', '10\u201315 min'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph(
    'Shooting with the coach is strongly recommended \u2014 it doubles the reach by tapping '
    'into his audience. The coach\'s endorsement also adds social proof that resonates '
    'with the fighting community far more than a physio promoting alone.'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2014 Right Track Physiotherapy & Performance Centre \u2014')
run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
run.font.size = Pt(10)
run.italic = True

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Seminar-Promo-Video-Scenarios.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
